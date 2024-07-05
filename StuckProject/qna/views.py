from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from .models import *
from accounts.models import *
import os
from django.contrib import messages
from django.db import IntegrityError


# openai
import openai
from django.conf import settings

# pdf에서 text 추출 
import fitz

# 이미지에서 text 추출
from google.cloud import vision
os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = 'service_account.json'

# 생성한 퀴즈 pdf 저장
from django.http import HttpResponse
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer


# 생성한 퀴즈 word 저장
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

@login_required
def landing_page(request):
    return render(request, 'quiz/landing.html')

# 퀴즈 생성 전 폴더 선택
@login_required
def select_folder(request, folder_id=None):
    if request.method == "POST":
        return redirect('qna:create-question-room', folder_id)
    if folder_id:
        folder = get_object_or_404(Folder, id=folder_id)
        children = folder.get_children() # 하위에 있는 모든 폴더
        path = folder.get_path() + "/"
    else:
        folder = None
        children = Folder.objects.filter(parent=None, user=request.user) # 루트에 있는 모든 폴더
        path = ""

    return render(request, 'qna/select_folder.html', {'folder': folder, 'children': children, 'path': path})


# 폴더 추가
def add_folder(request, parent_id):
    if parent_id == 0:
        parent = None
    else:
        parent = get_object_or_404(Folder, id=parent_id)

    if request.method == 'POST':
        name = request.POST['folder_name']

        try:
            Folder.objects.create(name=name, parent=parent, user=request.user)
            messages.success(request, f"{name} 폴더가 성공적으로 추가되었습니다.")
        except IntegrityError:
            messages.error(request, f"{name} 폴더 이름이 중복됩니다. 다른 이름을 사용해주세요.")

    if parent == None:
        return redirect('qna:select-folder', 0)
    return redirect('qna:select-folder', parent.id)


# Q&A Room 생성
def create_question_room(request, folder_id):
    if request.method == "POST":
        file = request.FILES.get('file')
        title = request.POST['title']
        folder = get_object_or_404(Folder, id=folder_id)
        # user = get_object_or_404(User, id=request.user.id)
        #  custom_user = get_object_or_404(User, id=request.user)

        question_room = QuestionRoom.objects.create(
            file=file,
            title=title,
            user=request.user,
            folder=folder,
            memo=""
        )

        return redirect('qna:enter-question-room', folder_id, question_room.id)

    context = {
        'folder_id': folder_id
    }
    return render(request, "qna/create_question_room.html", context)


# pdf 파일 열기 및 텍스트 추출
def extract_text_from_pdf(file_path):
    document = fitz.open(file_path)
    text = ""
    for page_num in range(len(document)):
        page = document.load_page(page_num)
        text += page.get_text()
    return text


# 이미지 인식하여 텍스트 추출 (google vision)
def extract_text_from_image(image_file_path):
    client = vision.ImageAnnotatorClient()

    # 파일 경로에서 이미지 파일을 읽어와 바이트 스트림으로 변환
    with open(image_file_path, 'rb') as image_file:
        content = image_file.read()

    image = vision.Image(content=content)

    # 텍스트 인식
    response = client.text_detection(image=image)
    texts = response.text_annotations

    if response.error.message:
        raise Exception(f'{response.error.message}')

    return texts[0].description if texts else ''


# Q&A Room 입장
def enter_question_room(request, folder_id, question_room_id):
    question_room = get_object_or_404(QuestionRoom, id=question_room_id)

    # 최근 열어본 문서 추가
    content_type = 'questionroom'
    user = get_object_or_404(User, id=request.user.id)
    custom_user = get_object_or_404(CustomUser, user=user)

    # 기록에 이미 있다면 삭제 
    RecentDocument.objects.filter(content_type=content_type, object_id=question_room.id, user=custom_user).delete()

    RecentDocument.objects.get_or_create(
        content_type=content_type, 
        object_id=question_room.id, 
        user=custom_user
    )

    # 최근 열어본 문서가 5개를 초과하면 가장 오래된 항목을 제거
    recent_docs = custom_user.recent_documents_from_user.all()
    if recent_docs.count() > 5:
        recent_docs.last().delete()

    if request.method == "POST":
        user_question = request.POST['user_question']
        file_path = question_room.file.path

        # 업로드한 파일이 pdf라면 
        if question_room.file and question_room.file.name.lower().endswith('.pdf'):
            # PyMuPDF 라이브러리를 사용하여 PDF 파일에서 텍스트 추출
            text = extract_text_from_pdf(file_path)
        else:
            # google cloud vision을 사용하여 이미지에서 텍스트 추출
            text = extract_text_from_image(file_path)

        openai.api_key = settings.OPENAI_API_KEY

        print(text)
        print(type)

        # 모델 - GPT 4 Turbo 선택
        model = "gpt-4-turbo"

        query = f"""${text}내용이 있어.
                ${user_question}
                """
        
        messages = [
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": text},
            {"role": "user", "content": query}
        ]

        # ChatGPT API 호출하기
        response = openai.ChatCompletion.create(
            model=model,
            messages=messages,
            temperature=0.8
        )

        # 응답 출력
        response = response['choices'][0]['message']['content']

        Chat.objects.create(
            user_chat=user_question,
            gpt_chat=response,
            question_room=question_room
        )

    chats = question_room.chats.all()
    context = {
        'chats': chats,
        'folder_id': folder_id,
        'question_room': question_room
    }

    return render(request, 'qna/question_room.html', context)


# Q&A에 해당하는 메모 저장 
def save_question_memo(request, folder_id, question_room_id):
    question_room = get_object_or_404(QuestionRoom, id=question_room_id)

    memo = request.POST['memo']

    question_room.memo = memo
    question_room.save()

    return redirect('qna:enter-question-room', folder_id, question_room.id)


# 메모 pdf 로 저장
def save_memo_as_pdf(request, folder_id, question_room_id):
    question_room = get_object_or_404(QuestionRoom, id=question_room_id)
    memo = question_room.memo
    title = question_room.title

    # 응답 설정
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="note.pdf"'

    # PDF 생성
    pdfmetrics.registerFont(TTFont('NanumGothic', 'static/fonts/NanumGothic.ttf'))
    styles = getSampleStyleSheet()
    
    # 기존 Title 스타일 수정
    styles['Title'].fontName = 'NanumGothic'
    styles['Title'].fontSize = 24
    styles['Title'].leading = 28
    styles['Title'].spaceAfter = 20
    
    styles['Normal'].fontName = 'NanumGothic'

    buffer = []
    doc = SimpleDocTemplate(response, pagesize=letter)

    # 제목 추가
    title_paragraph = Paragraph(title, styles['Title'])
    buffer.append(title_paragraph)
    buffer.append(Spacer(1, 12))  # Add space after title

    # 메모 내용 추가
    content = Paragraph(memo, styles['Normal'])
    buffer.append(content)

    doc.build(buffer)

    return response


# 메모 word 로 저장
def save_memo_as_word(request, folder_id, question_room_id):
    question_room = get_object_or_404(QuestionRoom, id=question_room_id)
    memo = question_room.memo
    title = question_room.title

    # 새로운 Word 문서 생성
    doc = Document()

    # 한글 폰트 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = 'NanumGothic'
    r = font.element.rPr.rFonts
    r.set(qn('w:eastAsia'), 'NanumGothic')

    # 제목 추가
    title_paragraph = doc.add_heading(level=1)
    run = title_paragraph.add_run(title)
    run.font.name = 'NanumGothic'
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), 'NanumGothic')
    run.font.size = Pt(24)

    # 메모 내용 추가
    memo_paragraph = doc.add_paragraph()
    run = memo_paragraph.add_run(memo)
    run.font.name = 'NanumGothic'
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), 'NanumGothic')
    run.font.size = Pt(12)

    # 응답 설정
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="note.docx"'

    # Word 문서 저장
    doc.save(response)

    return response


# 드래그 앤 드랍으로 질문 삭제
def delete_question_room(request, question_room_id):
    question_room = get_object_or_404(QuestionRoom, id=question_room_id)
    folder_id = question_room.folder.id
    question_room.delete()
    return redirect('quiz:folder-view', folder_id)


 # 드래그 앤 드랍으로 폴더 삭제
def move_question_room(request, current_folder_id, question_room_id, move_folder_id):
    questoin_room = get_object_or_404(QuestionRoom, id=question_room_id)
    move_folder = get_object_or_404(Folder, id=move_folder_id)

    questoin_room.folder = move_folder
    questoin_room.save()

    return redirect('quiz:folder-view', folder_id=current_folder_id)


# 질문 방 스크랩
def add_scrap_question_room(request, folder_id, question_room_id):
    user = get_object_or_404(User, id=request.user.id)
    custom_user = get_object_or_404(CustomUser, user=user)

    question_room = get_object_or_404(QuestionRoom, id=question_room_id)
    
    ScrapQuestionRoom.objects.get_or_create(user=custom_user, question_room=question_room)

    return redirect('qna:enter-question-room', folder_id, question_room.id)


# 질문 방 스크랩 취소
def remove_scrap_question_room(request, folder_id, question_room_id):
    user = get_object_or_404(User, id=request.user.id)
    custom_user = get_object_or_404(CustomUser, user=user)

    question_room = get_object_or_404(QuestionRoom, id=question_room_id)
    
    scrap_question_room = ScrapQuestionRoom.objects.get(user=custom_user, question_room=question_room)
    scrap_question_room.delete()

    return redirect('qna:enter-question-room', folder_id, question_room.id)


# 채팅 pdf 저장 
def save_chat_as_pdf(request, folder_id, question_room_id):
    question_room = get_object_or_404(QuestionRoom, id=question_room_id)
    title = question_room.title
    chat_messages = Chat.objects.filter(question_room=question_room).order_by('created_at')

    # 응답 설정
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="STUCK_CHAT.pdf"'

    # PDF 생성
    pdfmetrics.registerFont(TTFont('NanumGothic', 'static/fonts/NanumGothic.ttf'))
    styles = getSampleStyleSheet()

    # 기존 Title 스타일 수정
    styles.add(ParagraphStyle(name='Chat', fontName='NanumGothic', fontSize=12, leading=14))
    styles.add(ParagraphStyle(name='Date', fontName='NanumGothic', fontSize=12, leading=14, fontWeight='bold'))
    styles['Title'].fontName = 'NanumGothic'
    styles['Title'].fontSize = 24
    styles['Title'].leading = 28
    styles['Title'].spaceAfter = 20

    buffer = []
    doc = SimpleDocTemplate(response, pagesize=letter)

    # 제목 추가
    title_paragraph = Paragraph(title, styles['Title'])
    buffer.append(title_paragraph)
    buffer.append(Spacer(1, 12))  

    # 채팅 내용 추가
    for msg in chat_messages:
        date_text = f"{msg.created_at.strftime('%Y-%m-%d %H:%M:%S')}"
        user_chat_text = f"YOU: {msg.user_chat}"
        gpt_chat_text = f"AI: {msg.gpt_chat}"
        
        buffer.append(Paragraph(date_text, styles['Date']))
        buffer.append(Spacer(1, 6)) 
        buffer.append(Paragraph(user_chat_text, styles['Chat']))
        buffer.append(Spacer(1, 6)) 
        buffer.append(Paragraph(gpt_chat_text, styles['Chat']))
        buffer.append(Spacer(1, 24)) 

    doc.build(buffer)

    return response


# 채팅 워드로 저장
def save_chat_as_word(request, folder_id, question_room_id):
    question_room = get_object_or_404(QuestionRoom, id=question_room_id)
    title = question_room.title
    chat_messages = Chat.objects.filter(question_room=question_room).order_by('created_at')

    # 문서 생성
    doc = Document()

    # 제목
    doc.add_heading(title, level=1)

    # 채팅 저장
    for msg in chat_messages:
        date_text = f"{msg.created_at.strftime('%Y-%m-%d %H:%M:%S')}"
        user_chat_text = f"YOU: {msg.user_chat}"
        gpt_chat_text = f"AI: {msg.gpt_chat}"
        
        doc.add_paragraph(date_text, style='BodyText').bold = True
        doc.add_paragraph(user_chat_text, style='BodyText')
        doc.add_paragraph(gpt_chat_text, style='BodyText')
        doc.add_paragraph()  # Add space between message pairs

    # 응답 설정
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="STUCK_CHAT.docx"'

    doc.save(response)

    return response