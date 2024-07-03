from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from .models import *
from accounts.models import *
from todo.models import Routine, ToDo
import os
from django.contrib import messages
from django.db import IntegrityError
from datetime import date, timedelta

# openai
import openai
from django.conf import settings

# pdf에서 text 추출 
import fitz

# 이미지에서 text 추출
from google.cloud import vision
os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = 'service_account.json'

# 생성한 퀴즈 pdf 저장
from django.http import HttpResponse
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.pdfmetrics import stringWidth

# 생성한 퀴즈 word 저장
from docx import Document

# 메인페이지
def home(request, week_offset=0):
    if not request.user.is_authenticated:
        return render(request, 'quiz/home.html')

    user = get_object_or_404(User, id=request.user.id)
    custom_user = get_object_or_404(CustomUser, user=user)

    # 스크랩 정보
    folder_scraps = ScrapFolder.objects.filter(user=custom_user)
    quiz_scraps = ScrapQuiz.objects.filter(user=custom_user)
    question_room_scraps = ScrapQuestionRoom.objects.filter(user=custom_user)

    # 최근 열어본 문서
    recent_docs = custom_user.recent_documents_from_user.all()[:5]

    documents = []
    for doc in recent_docs:
        if doc.content_type == 'quiz':
            obj = get_object_or_404(Quiz, id=doc.object_id)
            folder_id = obj.folder.id
            documents.append({
                'type': 'quiz',
                'obj': obj,
                'folder_id': folder_id
            })
        elif doc.content_type == 'questionroom':
            obj = get_object_or_404(QuestionRoom, id=doc.object_id)
            folder_id = obj.folder.id
            documents.append({
                'type': 'questionroom',
                'obj': obj,
                'folder_id': folder_id
            })

    # 루틴 및 투두 정보
    routines = Routine.objects.filter(user=custom_user)
    todos = ToDo.objects.filter(routine__in=routines)

    # offset 만큼 week 변화
    today = date.today() + timedelta(weeks=week_offset)
    # 일요일부터 시작
    start_of_week = today - timedelta(days=(today.weekday() + 1) % 7)  # Sunday
    end_of_week = start_of_week + timedelta(days=6)

    week_days = []
    for i in range(7):
        day = start_of_week + timedelta(days=i)
        day_name = day.strftime('%a')
        day_todos = todos.filter(date=day)
        completed_count = day_todos.filter(completed=True).count()
        pending_count = day_todos.filter(completed=False).count()
        color = "blue" if completed_count > 0 else "gray"
        week_days.append({
            'day_name': day_name,
            'date': day.day,
            'todo_count': pending_count,
            'completed_count': completed_count,
            'color': color
        })

    # 주간에 포함된 월 이름 결정
    start_year = start_of_week.year
    end_year = end_of_week.year
    start_month = start_of_week.month
    end_month = end_of_week.month

    if start_year == end_year:
        if start_month == end_month:
            week_month = f"{start_year}년 {start_month}월"
        else:
            week_month = f"{start_year}년 {start_month}월 - {end_month}월"
    else:
        week_month = f"{start_year}년 {start_month}월 - {end_year}년 {end_month}월"

    # 월간 목표 달성률 계산
    start_of_month = date(today.year, today.month, 1)
    if today.month == 12:
        end_of_month = date(today.year, 12, 31)
    else:
        end_of_month = date(today.year, today.month + 1, 1) - timedelta(days=1)
    
    current_day = start_of_month
    month_days = []
    month_completed_count = 0
    month_total_count = 0
    while current_day <= end_of_month:
        day_name = current_day.strftime('%a')
        day_todos = todos.filter(date=current_day)
        completed_count = day_todos.filter(completed=True).count()
        pending_count = day_todos.filter(completed=False).count()

        color = "blue" if completed_count > 0 else "gray"

        month_days.append({
            'day_name': day_name,
            'date': current_day.day,
            'todo_count': pending_count,
            'completed_count': completed_count,
            'color': color
        })
        
        month_completed_count += completed_count
        month_total_count += completed_count + pending_count
        current_day += timedelta(days=1)

    month_completion_rate = (month_completed_count / month_total_count * 100) if month_total_count > 0 else 0

    # 주간 목표 달성률 계산
    week_completed_count = sum(day['completed_count'] for day in week_days)
    week_total_count = sum(day['completed_count'] + day['todo_count'] for day in week_days)
    week_completion_rate = (week_completed_count / week_total_count * 100) if week_total_count > 0 else 0

    # 일간 목표 달성률 계산
    day_todos = todos.filter(date=today)
    day_completed_count = day_todos.filter(completed=True).count()
    day_total_count = day_todos.count()
    day_completion_rate = (day_completed_count / day_total_count * 100) if day_total_count > 0 else 0

    context = {
        'folder_scraps': folder_scraps,
        'quiz_scraps': quiz_scraps,
        'question_room_scraps': question_room_scraps,
        'documents': documents,
        'week_days': week_days,
        'month_days': month_days,
        'week_offset': week_offset,
        'month_completion_rate': int(month_completion_rate),
        'week_completion_rate': int(week_completion_rate),
        'day_completion_rate': int(day_completion_rate),
        'week_month': week_month, 
    }

    return render(request, 'quiz/home.html', context)

# 폴더 조회
@login_required
def view_folder(request, folder_id=None):
    if folder_id:
        folder = get_object_or_404(Folder, id=folder_id)
        children = folder.get_children() # 하위에 있는 모든 폴더
        path = "Stuck/" + folder.get_path()
    else:
        folder = None
        children = Folder.objects.filter(parent=None, user=request.user) # 루트에 있는 모든 폴더
        path = ""

    user = get_object_or_404(User, id=request.user.id)
    custom_user = get_object_or_404(CustomUser, user=user)

    folder_scraps = ScrapFolder.objects.filter(user=custom_user)
    quiz_scraps = ScrapQuiz.objects.filter(user=custom_user)
    question_room_scraps = ScrapQuestionRoom.objects.filter(user=custom_user)

    print(quiz_scraps)
    context = {
        'folder': folder,
        'children': children,
        'path': path, 
        'folder_scraps': folder_scraps,
        'quiz_scraps': quiz_scraps,
        'question_room_scraps': question_room_scraps
    }

    return render(request, 'quiz/view_folder.html', context)


# 퀴즈 생성 전 폴더 선택
@login_required
def select_folder(request, folder_id=None):
    if request.method == "POST":
        return redirect('quiz:create-quiz', folder_id)
    if folder_id:
        folder = get_object_or_404(Folder, id=folder_id)
        children = folder.get_children() # 하위에 있는 모든 폴더
        path = "Stuck/" + folder.get_path()
    else:
        folder = None
        children = Folder.objects.filter(parent=None, user=request.user) # 루트에 있는 모든 폴더
        path = ""

    return render(request, 'quiz/select_folder.html', {'folder': folder, 'children': children, 'path': path})


# 폴더 추가
def add_folder(request, parent_id):
    if parent_id == 0:
        parent = None
    else:
        parent = get_object_or_404(Folder, id=parent_id)

    if request.method == 'POST':
        name = request.POST['folder_name']

        try:
            Folder.objects.create(name=name, parent=parent, user=request.user)
            messages.success(request, f"{name} 폴더가 성공적으로 추가되었습니다.")
        except IntegrityError:
            messages.error(request, f"{name} 폴더 이름이 중복됩니다. 다른 이름을 사용해주세요.")

    if parent is None:
        return redirect('quiz:select-folder', 0)
    return redirect('quiz:select-folder', parent.id)


# 폴더 드래그로 이동
def move_folder(request, folder_id, parent_id):
    folder = get_object_or_404(Folder, id=folder_id)

    if parent_id == 0:
        new_parent = None
    else:
        new_parent = get_object_or_404(Folder, id=parent_id)

    folder.parent = new_parent
    folder.save()

    if new_parent == None:
            return redirect('quiz:folder-view', 0)
    
    return redirect('quiz:folder-view', folder_id=new_parent.parent.id if new_parent.parent else 0)


# 퀴즈 드래그로 이동
def move_quiz(request, current_folder_id, quiz_id, move_folder_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)
    move_folder = get_object_or_404(Folder, id=move_folder_id)

    quiz.folder = move_folder
    quiz.save()

    return redirect('quiz:folder-view', folder_id=current_folder_id)


# pdf 파일 열기 및 텍스트 추출
def extract_text_from_pdf(file_path):
    document = fitz.open(file_path)
    text = ""
    for page_num in range(len(document)):
        page = document.load_page(page_num)
        text += page.get_text()
    return text


# 문제와 정답 항목 분리
def extract_questions_and_answers(response, type):
    print("all" + response)
    questions_part, answers_part = response.split("Answers:")
    
    # 문제 부분을 줄 단위로 분리
    questions_lines = questions_part.split("\n\n")
    
    # 문제를 담을 리스트 초기화
    questions = []
    
    # 각 질문을 확인하며 리스트에 추가
    for question in questions_lines:
        if question.strip():
            questions.append(question.strip())
    
    # 답안 부분을 줄 단위로 분리
    answers = answers_part.strip().split('\n')
    
    # 객관식일 경우
    if type == "객관식":
        # 숫자 이외의 값 제거
        answers = [answer.strip() for answer in answers if answer.strip().isdigit()]
    
    return questions, answers


# 이미지 인식하여 텍스트 추출 (google vision)
def extract_text_from_image(image_file_path):
    client = vision.ImageAnnotatorClient()

    # 파일 경로에서 이미지 파일을 읽어와 바이트 스트림으로 변환
    with open(image_file_path, 'rb') as image_file:
        content = image_file.read()

    image = vision.Image(content=content)

    # 텍스트 인식
    response = client.text_detection(image=image)
    texts = response.text_annotations

    if response.error.message:
        raise Exception(f'{response.error.message}')

    return texts[0].description if texts else ''


# 퀴즈 생성 
def create_question(request, folder_id):
    # print(settings.OPENAI_API_KEY)
    if request.method == "GET":
        return render(request, 'quiz/create_quiz.html', {'folder_id':folder_id})
    
    folder = get_object_or_404(Folder, id=folder_id)

    file = request.FILES.get('file')
    title = request.POST['title']
    question_num = request.POST['question_num']
    type = request.POST['type']

    # 퀴즈 생성
    quiz = Quiz.objects.create(
        folder=folder,
        file=file,
        title=title,
        question_num=int(question_num),
        type=type
    )

    file_path = quiz.file.path

    # 업로드한 파일이 pdf라면 
    if quiz.file and quiz.file.name.lower().endswith('.pdf'):
        # PyMuPDF 라이브러리를 사용하여 PDF 파일에서 텍스트 추출
        text = extract_text_from_pdf(file_path)
    else:
        # google cloud vision을 사용하여 이미지에서 텍스트 추출
        text = extract_text_from_image(file_path)

    openai.api_key = settings.OPENAI_API_KEY

    print(text)
    print(type)

    # 모델 - GPT 4 Turbo 선택
    model = "gpt-4-turbo"

    # 주관식을 선택했을 경우 프롬프트
    if type == "주관식":
        query = f"""
            이 내용을 기반으로 {question_num}개의 {type} 문제를 내줘. 
            문제의 답안은 서술형이 아닌 단어 하나여야 해.
            각 문제의 시작은 문제 1. 처럼 문제 + 문제 번호 수 .이야. 
            문제는 꼭 물음표로 끝나야 해.
            문제는 모두 \n\n\n으로 꼭 구분해줘.
            예를 들어서 문제 1. 첫번째 문제 제시.\n\n\n문제2. 두번째 문제 제시.\n\n\n처럼 문제를 \n\n\n으로 구분해서 응답해줘.
            문제를 모두 내준 후 답안은 Answers: 으로 구분해서 아래에 답도 알려줘. Answers: 외에 구분선을 넣으면 안돼.
            Answers: 아래에 답안을 알려줄때는 문제 번호, 문항 내용 없이 문제 답만 작성하고 답 구분은 \n으로 해줘.
            꼭 답안이 서술형이 아닌 단어 하나로 구성해줘.
        """
    # 객관식을 선택했을 경우 프롬프트
    else : 
        query = f"""
        이 내용을 기반으로 {question_num}개의 {type} 문제를 내줘. 
        각 문제의 시작은 문제 1. 처럼 문제 + 문제 번호 수 .이야. 
        문제는 꼭 물음표로 끝나야 해.
        각 문제의 항목은 1. 2. 3. 4. 처럼 [숫자.]으로 시작하고 문제마다 4개야. 
        문제는 모두 \n\n\n으로 구분해줘. 
        문제를 모두 내준 후 답안은 Answers: 으로 구분해서 아래에 답도 알려줘. Answers: 외에 구분선을 넣으면 안돼.
        Answers: 아래에 답안을 알려줄때는 문제 번호, 문항 내용 없이 문제 답 번호만 작성하고 답 구분은 \n으로 해줘.
        예를 들어서 1\n2\n4\n 처럼."""

    # 메시지 설정하기
    messages = [
        {"role": "system", "content": "You are a helpful assistant."},
        {"role": "user", "content": text},
        {"role": "user", "content": query}
    ]

    # ChatGPT API 호출하기
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0.8
    )

    # 응답 출력
    response = response['choices'][0]['message']['content']

    # 문제와 답안 추출
    questions, answers = extract_questions_and_answers(response, type)

    # 테스트 출력
    print("Questions:")
    for q in questions:
        print(q)
        print()

    print("Answers:")
    for a in answers:
        print(a)
        print()

    print(f"questions len {len(questions)}")
    print(f"answers len {len(answers)}")
    
    # 각 질문과 답안을 개별적으로 저장
    for i in range(0, int(question_num)):
        Question.objects.create(
            quiz=quiz,
            ai_question=questions[i],
            correct_answer=answers[i]
        )

    return redirect('quiz:test', folder_id, quiz.id)


# 퀴즈 풀기
def test(request, folder_id, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)
    questions = list(quiz.questions.all())
    
    if request.method == 'POST':
        for question in questions:
            user_answer = request.POST.get(f'answer_{question.id}')
            if user_answer is not None:
                question.user_answer = user_answer
                question.save()
        
        return redirect('quiz:quiz-results', folder_id=folder_id, quiz_id=quiz_id)
    
    message = ""
    if quiz.type == "주관식":
        message = "답안은 띄어쓰기 없이 작성해주세요."

    context = {
        'questions': questions,
        'quiz': quiz,
        'type': quiz.type,
        'message': message,
        'folder_id': folder_id
    }
    
    return render(request, 'quiz/test.html', context)


# 퀴즈 결과 확인
def quiz_results(request, folder_id, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)
    questions = quiz.questions.all()

    correct_count = sum(1 for q in questions if q.status) # 맞은 개수
    total_questions = questions.count() # 전체 문제 수

    quiz.score = correct_count
    quiz.save()

    context = {
        'quiz': quiz,
        'questions': questions,
        'correct_count': correct_count,
        'total_questions': total_questions,
        'folder_id':folder_id
    }
    
    return render(request, 'quiz/results.html', context)


# 생성된 문제 전체 확인
def view_questions(request, folder_id, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)
    questions = quiz.questions.all()

    correct_count = sum(1 for q in questions if q.status) # 맞은 개수
    total_questions = questions.count() # 전체 문제 수

    context = {
        "quiz": quiz,
        'questions': questions,
        'correct_count': correct_count,
        'total_questions': total_questions,
        'folder_id':folder_id
    }

    # 최근 열어본 문서 추가
    content_type = 'quiz'
    user = get_object_or_404(User, id=request.user.id)
    custom_user = get_object_or_404(CustomUser, user=user)

    # 기존 이미 있다면 삭제 
    RecentDocument.objects.filter(content_type=content_type, object_id=quiz.id, user=custom_user).delete()

    RecentDocument.objects.create(
        content_type=content_type, 
        object_id=quiz.id, 
        user=custom_user
    )

    # 최근 열어본 문서가 5개를 초과하면 가장 오래된 항목을 제거
    recent_docs = custom_user.recent_documents_from_user.all()
    if recent_docs.count() > 5:
        recent_docs.last().delete()

    return render(request, 'quiz/results.html', context)


# 생성된 문제 중 틀린 문제만 확인
def view_wrong_questions(request, folder_id, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)
    questions = quiz.questions.filter(status=False)

    correct_count = sum(1 for q in questions if q.status) # 맞은 개수
    total_questions = questions.count() # 전체 문제 수

    context = {
        'quiz': quiz,
        'questions': questions,
        'correct_count': correct_count,
        'total_questions': total_questions,
        'folder_id':folder_id
    }
    return render(request, 'quiz/results.html', context)


# PDF 너비에 따라 텍스트 줄바꿈 
def wrap_text(text, width, font_name, font_size):
    lines = text.split('\n')
    wrapped_lines = []
    for line in lines:
        words = line.split(' ') # 단어 단위로
        temp_line = ''

        for word in words:
            # 이전 word 합 + 새로 작성하고자 하는 word 가 텍스트의 너비보다 작은가? (텍스트, 폰트 적용하여 계산) 
            if stringWidth(temp_line + word + ' ', font_name, font_size) <= width:
                temp_line += word + ' '

            else:
                # 지정된 너비를 초과하면 줄바꿈 
                wrapped_lines.append(temp_line.strip())
                # 빈 문자 처리
                temp_line = word + ' ' 

        wrapped_lines.append(temp_line.strip())
    return wrapped_lines


# 질문과 선택 항목 분리
def format_question_text(text):
    # 모든 질문은 ?로 끝남 
    parts = text.split('?')

    # 객관식인가?
    if len(parts) > 1:
        # 질문 부분
        question_part = parts[0] + '?'
        # 선택 항목 부분
        options_part = parts[1].strip()
        # 공백을 기준으로 단어로 나눔 
        options = options_part.split(' ')

        formatted_options = []
        option_text = ''

        for option in options:
            # 각 단어가 숫자와 .으로 시작하는 경우 (ex 1.~~)
            if option[0].isdigit() and option[1] == '.':
                if option_text:
                    # 선택지에 추가
                    formatted_options.append(option_text.strip())
                option_text = option + ' '

            else:
                option_text += option + ' '

        if option_text:
            formatted_options.append(option_text.strip())
        
        # 선택지 리스트를 줄바꿈으로 연결
        formatted_text = question_part + '\n' + '\n'.join(formatted_options)
        return formatted_text
    
    # 주관식은 len이 1이므로 text만 반환 
    return text


# 생성한 퀴즈 pdf 저장 
def save_quiz_as_pdf(request, folder_id, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)

    # 퀴즈 전체 질문
    questions = quiz.questions.all()

    # 응답
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="Stuck_QUIZ.pdf"'

    # 나눔고딕 폰트
    pdfmetrics.registerFont(TTFont('NanumGothic', 'static/fonts/NanumGothic.ttf'))

    # PDF 캔버스 생성 
    p = canvas.Canvas(response, pagesize=letter)

    # 제목
    title_font_size = 16
    p.setFont("NanumGothic", title_font_size)
    title_text = f"퀴즈 : {quiz.title}"
    title_width = stringWidth(title_text, "NanumGothic", title_font_size)
    p.drawCentredString(letter[0] / 2, 750, title_text)
    p.setFont("NanumGothic", 12)

    # 퀴즈 정보
    p.drawString(90, 710, f"생성 날짜: {quiz.created_at}")
    p.drawString(90, 690, f"유형: {quiz.type}")
    p.drawString(90, 670, f"문제 수: {quiz.question_num}")

    y_position = 630 # 첫번째 장의 출력할 y축
    line_height = 15 # text 한 줄 높이
    max_width = 450  # 페이지 너비

    # 질문 출력
    for question in questions:
        question_text = f"{question.ai_question}"
        wrapped_lines = wrap_text(question_text, max_width, "NanumGothic", 12)
        print(wrapped_lines)
        for wrapped_line in wrapped_lines:
            if y_position < 50: 
                p.showPage()  
                p.setFont("NanumGothic", 12)
                y_position = 750  
            
            p.drawString(90, y_position, wrapped_line)
            y_position -= line_height
        
        y_position -= 20  

    p.showPage()
    p.setFont("NanumGothic", 12)
    y_position = 750  

    # 새로운 페이지에 답안 출력
    y_position -= line_height

    question_num = 1
    for question in questions:
        answer_text = f"{question_num}번 답안: {question.correct_answer}"
        wrapped_lines = wrap_text(answer_text, max_width, "NanumGothic", 12)
        for wrapped_line in wrapped_lines:
            if y_position < 50:
                p.showPage() 
                p.setFont("NanumGothic", 12)
                y_position = 750  
            
            p.drawString(90, y_position, wrapped_line)
            y_position -= line_height
        
        y_position -= 20  
        question_num += 1

    # PDF 저장 및 다운로드 진행 
    p.showPage()
    p.save()

    return response


# 생성한 퀴즈 word 저장 
def save_quiz_as_word(request, folder_id, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)

    # 퀴즈 전체 질문
    questions = quiz.questions.all()

    # Word 문서 생성
    document = Document()

    # 제목 추가
    document.add_heading(f'퀴즈 : {quiz.title}', level=1)

    # 퀴즈 정보 추가
    document.add_paragraph(f'생성 날짜: {quiz.created_at}')
    document.add_paragraph(f'유형: {quiz.type}')
    document.add_paragraph(f'문제 수: {quiz.question_num}')
    document.add_paragraph()

    # 질문 출력
    for question in questions:
        question_text = question.ai_question.split('\n')
        for line in question_text:
            document.add_paragraph(line)
        document.add_paragraph()  # 질문 사이의 간격 추가

    # 새로운 페이지에 답안 출력
    document.add_page_break()
    document.add_heading('답안', level=2)

    question_num = 1
    for question in questions:
        answer_text = f'{question_num}번 답안: {question.correct_answer}'
        document.add_paragraph(answer_text)
        question_num += 1

    # Word 문서 응답
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=quiz_report.docx'
    document.save(response)

    return response


# 폴더 스크랩
@login_required
def add_scrap_folder(request, folder_id):
    folder = get_object_or_404(Folder, id=folder_id)
    user = get_object_or_404(User, id=request.user.id)
    custom_user = get_object_or_404(CustomUser, user=user)

    ScrapFolder.objects.get_or_create(user=custom_user, folder=folder)
    print(request.user.customuser.scrap_folders)
    return redirect('quiz:folder-view', folder_id=folder.id)


# 폴더 스크랩 취소
@login_required
def remove_scrap_folder(request, folder_id):
    user = get_object_or_404(User, id=request.user.id)
    custom_user = get_object_or_404(CustomUser, user=user)
    folder = get_object_or_404(Folder, id=folder_id)
    scraps = ScrapFolder.objects.filter(user=custom_user, folder=folder)
    scraps.delete()
    return redirect('quiz:folder-view', folder_id)



# 퀴즈 스크랩
def add_scrap_quiz(request, folder_id, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)
    user = get_object_or_404(User, id=request.user.id)
    custom_user = get_object_or_404(CustomUser, user=user)
    
    ScrapQuiz.objects.get_or_create(user=custom_user, quiz=quiz)

    return redirect('quiz:view-questions', folder_id, quiz.id)


# 퀴즈 스크랩 취소
def remove_scrap_quiz(request, folder_id, quiz_id):
    user = get_object_or_404(User, id=request.user.id)
    custom_user = get_object_or_404(CustomUser, user=user)
    quiz = get_object_or_404(Quiz, id=quiz_id)
    
    scrap_quiz = ScrapQuiz.objects.get(user=custom_user, quiz=quiz)
    scrap_quiz.delete()
    
    return redirect('quiz:view-questions', folder_id, quiz.id)


# 드래그 앤 드랍으로 폴더 삭제
def delete_folder(request, folder_id):
    folder = get_object_or_404(Folder, id=folder_id)
    parent = folder.parent
    folder.delete()
    if(parent == None):
        return redirect('quiz:folder-view', 0)
    return redirect('quiz:folder-view', parent.id)


# 드래그 앤 드랍으로 퀴즈 삭제
def delete_quiz(request, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)
    folder_id = quiz.folder.id
    quiz.delete()
    return redirect('quiz:folder-view', folder_id)


# 폴더 검색 
def search_folder(request, folder_id):
    search_folder_name = request.GET.get('search_folder_name', '')

    folder = Folder.objects.filter(name=search_folder_name).first()

    if not folder:
            messages.error(request, "폴더가 존재하지 않습니다.")
            return redirect('quiz:folder-view', folder_id)
    messages.success(request, "폴더가 성공적으로 검색되었습니다.")
    return redirect('quiz:folder-view', folder.id)



# 목표 달성률
from todo.models import ToDo

def get_rate(request):
    print("rate has been called")

    todos = ToDo.objects.all()
    total_todos = todos.count()
    completed_todos = todos.filter(completed=True).count()
    completion_rate = (completed_todos / total_todos * 100) if total_todos > 0 else 0

    print(f"Completion Rate: {completion_rate}")

    return render(request, 'quiz/home.html', {'completion_rate': completion_rate})
